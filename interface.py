import streamlit as st, os, json, datetime, io
from docx import Document

def init_ui_styles():
    st.markdown("""<style>
        .stTextArea textarea { border: 1px solid #ddd!important; background:#fafafa!important; }
        .header-style { font-size: 1.6rem; font-weight: 700; color: #1d1d1f; margin-bottom: 10px; }
        .stExpander { border: 1px solid #eaeaea!important; border-radius: 8px!important; margin-bottom: 12px!important; }
        footer {visibility: hidden;}
        .sidebar .stButton button { margin-bottom: 2px; text-align: left; border: none !important; }
        button[key^="pre_"], button[key^="d_"], button[key^="sd_"], button[key^="f_"] { color: #ff4b4b !important; }
    </style>""", unsafe_allow_html=True)

def get_docx():
    doc = Document(); doc.add_heading(st.session_state.inv_title, 0)
    for t, k in [("Methods","m_c"),("Results","r_c"),("Discussion","d_c")]:
        doc.add_heading(t, 1); doc.add_paragraph(st.session_state[k])
    b = io.BytesIO(); doc.save(b); return b.getvalue()

def render_sidebar(D, df_schema, save, undo, snap):
    with st.sidebar:
        st.header("Projects")
        c1, c2 = st.columns(2)
        if c1.button("➕ New", use_container_width=True):
            for k, v in df_schema.items(): st.session_state[k] = v
            st.rerun()
        if c2.button("↩️ Undo", use_container_width=True): undo()
        st.divider()
        
        p_files = sorted([f for f in os.listdir(D) if f.endswith(".json")])
        for f in p_files:
            p_n = f.replace(".json","").replace("_"," ")
            col1, col2 = st.columns([0.8, 0.2])
            if col1.button(p_n, key=f"l_{f}", use_container_width=True):
                try:
                    with open(os.path.join(D, f), "r") as pf: st.session_state.update(json.load(pf))
                    st.rerun()
                except:
                    st.error(f"'{p_n}' corrupted.")
                    if st.button("Force Clear", key=f"f_{f}"): os.remove(os.path.join(D, f)); st.rerun()
            with col2:
                if st.session_state.get(f"c_{f}", False):
                    if st.button("!!", key=f"d_{f}"): os.remove(os.path.join(D, f)); st.session_state[f"c_{f}"]=False; st.rerun()
                elif st.button("X", key=f"pre_{f}"): st.session_state[f"c_{f}"] = True; st.rerun()
        st.divider()
        with st.expander("⏳ Timeline"):
            for i, v in enumerate(reversed(st.session_state.v_h)):
                sc1, sc2 = st.columns([0.8, 0.2])
                if sc1.button(v['ts'], key=f"sr_{i}"):
                    for k, val in v['db'].items(): st.session_state[k] = val
                    st.rerun()
                if sc2.button("X", key=f"sd_{i}"): st.session_state.v_h.pop(len(st.session_state.v_h)-1-i); save(); st.rerun()
        st.download_button("📝 Export", data=get_docx(), file_name=f"{st.session_state.inv_title}.docx", use_container_width=True)